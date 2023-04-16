from sys import argv
from docx import Document


inputfile = str(argv[1])+str(argv[2])+".vtt"
timestampfile = str(argv[1])+"timestamps.docx"
transcriptionfile = str(argv[1])+"transcription.docx"


# Open inputfile and create timestampfile & transcriptionfile
inputf = open(inputfile,"r")
timestamps = Document()
transcription = Document()


# Sort the lines from the inputfile into timestampfile & transcriptionfile
linecount = 1
for line in inputf:
	if linecount%3==0:
		timestamps.add_para(line)
	elif (linecount-1)%3==0:
		if linecount != 1: # Skip the first line "WEBVTT"
			transcription.add_para(line)
	linecount += 1


# Check that the number of timestamps matches the number of lines pre-translation
if len(timestamps.paragraphs) == len(transcription.paragraphs):
	print("Pre-translation check passed: Number of timestamps and number of lines match")
else:
	print("WARNING! Pre-translation check failed: Number of timestamps and number of lines do NOT match")
	print("Number of timestamps: "+str(len(timestamps.paragraphs)))
	print("Number of lines: "+str(len(transcription.paragraphs)))


# Close inputfile and save timestampfile & transcriptionfile
inputf.close()
timestamps.save(timestampfile)
transcription.save(transcriptionfile)	
