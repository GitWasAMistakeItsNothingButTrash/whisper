from docx import Document

# Define filepaths
print ("Please enter the absolute path to Whisper's speech-to-text transcription (including timestamps): ")
inputfile = raw_input("")

print ("Please enter the absolute path at which to save the raw transcription without timestamps: ")
transcriptionfile = raw_input("")

print ("Please enter the absolute path at which to save the raw translation without timestamps: ")
translationtionfile = raw_input("")

print ("Please enter the absolute path at which to save the translated subtitles (including timestamps): ")
outputfile = raw_input("")


inputf = open(inputfile,"r")
transcription = docx.Document()

timestamps = []

linecount = 1

for line in inputf:
	if linecount%3==0:
		timestamps.append(line)
	elif (linecount-1)%3==0:
		if linecount != 0:
			transcription.add_para(line)
	linecount += 1

# Check that the number of timestamps matches the number of lines pre-translation
if len(timestamps) == len(transcription.paragraphs):
	print("Pre-translation check passed: Number of timestamps and number of lines match")
elif len(timestamps) > len(transcription.paragraphs):
	print("WARNING! Pre-translation check failed: Number of timestamps exceeds number of lines")
elif len(timestamps) < len(transcription.paragraphs):
	print("WARNING! Pre-translation check failed: Number of lines exceeds number of timestamps")	

inputf.close()
transcription.save(transcriptionfile)	

!!! Translate transcriptionfile --> translationfile

# Open translation and create output file
translation = Document(translationfile)
output = open(outputfile,"w")

# Write format header and empty line to output
output.write("WEBVTT")
output.write("\n")

# Check that the number of timestamps matches the number of lines post-translation
if len(timestamps) == len(translation.paragraphs):
	print("Post-translation check passed: Number of timestamps and number of lines match")
elif len(timestamps) > len(translation.paragraphs):
	print("WARNING! Post-translation check failed: Number of timestamps exceeds number of lines")
elif len(timestamps) < len(translation.paragraphs):
	print("WARNING! Post-translation check failed: Number of lines exceeds number of timestamps")

# Merge timestamps and translation, including empty lines
for line in range(len(timestamps)):
	output.write(str(timestamps[line]))
	output.write(str(translation.paragraphs[line].text))
	output.write("\n")

# Save and close output file
output.close()
