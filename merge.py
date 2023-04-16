from sys import argv
from docx import Document


outputfile = str(argv[1])+str(argv[2])+".vtt"
timestampfile = str(argv[1])+"timestamps.docx"
translationfile = str(argv[1])+"translation.docx"


# Create outputfile and open timestampfile & translationfile
output = open(outputfile,"w")
timestamps = Document(timestampfile)
translation = Document(translationfile)

# Write format header and empty line to output
output.write("WEBVTT")
output.write("\n")


# Check that the number of timestamps matches the number of lines post-translation
if len(timestamps.paragraphs) == len(translation.paragraphs):
	print("Post-translation check passed: Number of timestamps and number of lines match")
else:
	print("WARNING! Post-translation check failed: Number of timestamps and number of lines do NOT match")
	print("Number of timestamps: "+str(len(timestamps.paragraphs)))
	print("Number of lines: "+str(len(translation.paragraphs)))


# Merge timestamps and translation, including empty lines
for line in range(len(timestamps)):
	output.write(str(timestamps.paragraphs[line].text))
	output.write(str(translation.paragraphs[line].text))
	output.write("\n")

# Save and close output file
output.close()
